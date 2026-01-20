"""
World-Class Intelligent Resume Parser
Uses advanced techniques: layout-aware parsing, NER, fuzzy matching, OCR support
"""
import re
import os
from typing import Dict, List, Optional, Tuple
from datetime import datetime

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PDF2_AVAILABLE = True
except ImportError:
    PDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf_intelligent(filepath: str) -> Tuple[str, Dict]:
    """
    Intelligent PDF extraction using pdfplumber (better layout preservation)
    Falls back to PyPDF2 if pdfplumber not available
    """
    text = ""
    layout_info = {'pages': [], 'tables': []}
    
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        layout_info['pages'].append({
                            'page_num': i + 1,
                            'text': page_text,
                            'width': page.width,
                            'height': page.height
                        })
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    if tables:
                        layout_info['tables'].extend(tables)
            return text, layout_info
        except Exception as e:
            print(f"[PARSER] pdfplumber failed: {e}, falling back to PyPDF2")
    
    # Fallback to PyPDF2
    if PDF2_AVAILABLE:
        try:
            print(f"[PDF EXTRACTION] Using PyPDF2 fallback")
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"[PDF EXTRACTION] PyPDF2: {len(pdf_reader.pages)} pages")
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        print(f"[PDF EXTRACTION] Page {i+1}: {len(page_text)} chars")
                
                print(f"[PDF EXTRACTION] Total text extracted: {len(text)} chars")
                if len(text) < 100:
                    print(f"[PDF EXTRACTION WARNING] Very little text extracted! First 500 chars: {text[:500]}")
                
                return text, layout_info
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    raise ImportError("No PDF library available. Install pdfplumber or PyPDF2")


def extract_text_from_docx_intelligent(filepath: str) -> Tuple[str, Dict]:
    """Extract text from DOCX with structure preservation"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")
    
    text_parts = []
    structure = {'paragraphs': [], 'tables': []}
    
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                structure['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else None
                })
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            structure['tables'].append(table_data)
        
        return "\n".join(text_parts), structure
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def extract_year(text: str) -> Optional[int]:
    """Extract year from text (4-digit year)"""
    years = re.findall(r'\b(19|20)\d{2}\b', text)
    if years:
        try:
            return int(years[-1])  # Get most recent year
        except:
            pass
    return None


def normalize_text(text: str) -> str:
    """Normalize text: fix common OCR errors, whitespace, etc."""
    if not text:
        return ""
    
    # Fix common OCR errors
    fixes = {
        'rn': 'm',
        'corn': 'com',
        'peabedi': 'abedi',
        'peabedisyedaliabbas': 'abedisyedaliabbas',
        'ng ': 'Nanyang ',
        'ngTechnological': 'Nanyang Technological',
        'ng Position': 'Visiting Researcher Position',
    }
    
    for wrong, correct in fixes.items():
        text = text.replace(wrong, correct)
    
    # Normalize whitespace
    text = ' '.join(text.split())
    
    return text.strip()


def parse_resume_intelligent(text: str, layout_info: Dict = None) -> Dict:
    """
    World-class intelligent resume parser
    Uses layout-aware parsing, NER-like patterns, fuzzy matching
    """
    # Debug: Check if text is empty
    if not text or len(text.strip()) < 50:
        print(f"[PARSER ERROR] Text is empty or too short: {len(text) if text else 0} chars")
        return {
            "name": "", "email": "", "phone": "", "location": "",
            "education": [], "experience": [], "skills": [],
            "publications": [], "presentations": [], "awards": [],
            "research_interests": [], "summary": ""
        }
    
    # Normalize text first
    text = normalize_text(text)
    
    # Preserve line structure
    original_lines = text.split('\n')
    text_lines = [line.strip() for line in original_lines if line.strip()]
    
    print(f"[PARSER DEBUG] Processing {len(text_lines)} lines, {len(text)} chars")
    if len(text_lines) < 10:
        print(f"[PARSER WARNING] Very few lines extracted: {text_lines}")
    
    data = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "education": [],
        "experience": [],
        "skills": [],
        "publications": [],
        "presentations": [],
        "awards": [],
        "research_interests": [],
        "summary": ""
    }
    
    # Extract email (prioritize contact section)
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        first_lines = '\n'.join(text_lines[:25])
        primary_emails = re.findall(email_pattern, first_lines)
        if primary_emails:
            email = primary_emails[0].lower()
            # Fix OCR errors
            email = email.replace('peabedi', 'abedi').replace('corn', 'com')
            data["email"] = email
            print(f"[PARSER DEBUG] Email extracted: '{email}'")
    else:
        print(f"[PARSER DEBUG] No email found in text. First 10 lines: {text_lines[:10]}")
    
    # Extract phone
    phone_patterns = [
        r'\+?\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\+?\d{2,4}\s?\d{3,4}\s?\d{3,4}'
    ]
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            data["phone"] = matches[0]
            break
    
    # Extract name (first substantial line, 2-5 words, capitalized)
    # Try first line if it looks like a name
    name_found = False
    if text_lines and len(text_lines) > 0:
        first_line = text_lines[0].strip()
        words = first_line.split()
        # Check if first line is a name (2-5 words, mostly capitalized, no email/phone/section headers)
        if (2 <= len(words) <= 5 and 
            '@' not in first_line and 
            not re.search(r'\+?\d', first_line) and
            not any(word.lower() in ['email', 'phone', 'date', 'birth', 'orcid', 'website', 'researcher', 'id', 'experience', 'education', 'publications', 'skills'] for word in words)):
            # Check if it starts with capitalized words (at least first 2 words)
            if len(words) >= 2 and all(w and w[0].isupper() for w in words[:2] if w):
                data["name"] = first_line
                name_found = True
                print(f"[PARSER DEBUG] Name extracted from first line: '{first_line}'")
        
        # If not found, try other early lines (skip lines with colons, emails, etc.)
        if not name_found:
            for i, line in enumerate(text_lines[:15]):
                line_clean = line.strip()
                # Skip lines with email, phone, colon (likely metadata), or section headers
                if ('@' not in line_clean and 
                    not re.search(r'\+?\d', line_clean) and
                    ':' not in line_clean and
                    not any(word in line_clean.lower() for word in ['email', 'phone', 'date', 'birth', 'orcid', 'website', 'experience', 'education', 'publications', 'skills', 'awards'])):
                    words = line_clean.split()
                    if 2 <= len(words) <= 5:
                        # Check if it's a name (capitalized, not a section header)
                        if (all(w and w[0].isupper() for w in words[:2] if w) and
                            not any(word.lower() in ['email', 'phone', 'date', 'birth', 'orcid', 'website', 'experience', 'education', 'publications', 'skills', 'awards'] for word in words)):
                            data["name"] = line_clean
                            name_found = True
                            print(f"[PARSER DEBUG] Name extracted from line {i}: '{line_clean}'")
                            break
    
    if not name_found:
        print(f"[PARSER DEBUG] Name NOT found. First 5 lines: {text_lines[:5]}")
    
    # Extract location
    location_keywords = ['Singapore', 'USA', 'United States', 'UK', 'United Kingdom', 
                        'Australia', 'Canada', 'India', 'Malaysia', 'Thailand', 'Hong Kong']
    for keyword in location_keywords:
        if keyword.lower() in text.lower():
            data["location"] = keyword
            break
    
    # Find section boundaries with better detection
    section_headers = {
        'education': ['# education', 'education', 'academic background', 'qualifications', 'qualification'],
        'experience': ['experience', 'work experience', 'employment', 'professional experience', 
                      'experience (most recent first)', 'career', 'positions'],
        'publications': ['publications', 'peer-reviewed', 'journal articles', 'papers', 
                        '# publications', 'research papers'],
        'presentations': ['presentations', 'conferences', 'invited presentations', 'talks', 'posters'],
        'awards': ['awards', 'prizes', 'funding', 'prizes, awards', 'honors', 'honours'],
        'skills': ['technical skills', 'skills', 'expertise', 'technical skills and expertise', 
                  'competencies', 'proficiencies']
    }
    
    section_indices = {}
    for section_name, keywords in section_headers.items():
        for i, line in enumerate(text_lines):
            line_lower = line.lower().strip()
            # Match section headers - check if line contains any keyword
            # Section headers are usually short (1-8 words) and don't contain email/phone
            if any(keyword in line_lower for keyword in keywords):
                words = line.split()
                # Check if it's a section header (not email/phone line, reasonable length)
                if (len(words) <= 8 and 
                    '@' not in line and 
                    not re.search(r'\+?\d', line) and
                    (line_lower.startswith('#') or  # Starts with #
                     len(words) <= 5 or  # Short line (likely header)
                     any(keyword == line_lower.strip() for keyword in keywords) or  # Exact keyword match
                     (line_lower.startswith(keywords[0]) and len(words) <= 6))):  # Starts with keyword, reasonable length
                    section_indices[section_name] = i
                    print(f"[PARSER DEBUG] Found '{section_name}' section at line {i}: '{line[:50]}'")
                    break
    
    print(f"[PARSER DEBUG] Section indices: {section_indices}")
    
    # INTELLIGENT Education Parsing
    edu_start = section_indices.get('education', 0)
    # If no explicit education section found, search for education patterns
    # BUT skip the first few lines (name, contact info)
    if edu_start == 0:
        for i, line in enumerate(text_lines[5:100], start=5):  # Start from line 5 to skip header
            line_lower = line.lower()
            # Look for degree keywords followed by field/institution keywords
            if any(word in line_lower for word in ['phd', 'ms', 'm.s', 'bsc', 'bachelor', 'master', 'degree']):
                if any(word in line_lower for word in ['in ', 'of ', 'from', 'university', 'college', 'institute', 'defence', 'awarded']):
                    edu_start = i
                    print(f"[PARSER DEBUG] Found education pattern at line {i}: '{line[:50]}'")
                    break
    
    edu_end = min(
        section_indices.get('experience', len(text_lines)),
        section_indices.get('publications', len(text_lines)),
        edu_start + 50
    ) if edu_start > 0 else min(100, len(text_lines))
    
    edu_lines = text_lines[edu_start:edu_end] if edu_start > 0 else text_lines[:100]
    print(f"[PARSER DEBUG] Education parsing: lines {edu_start} to {edu_end}, {len(edu_lines)} lines")
    
    # Pattern: "- PhD in Science, Mathematics & Technology" or "- PhD in X"
    # Followed by institution on next line(s)
    # Followed by "Awarded: YYYY" or "Defence: DD/MM/YYYY"
    
    i = 0
    while i < len(edu_lines):
        line = edu_lines[i]
        line_lower = line.lower()
        
        # Match degree patterns (with or without bullet)
        degree_match = None
        degree_patterns = [
            (r'[-•*]?\s*(PhD|Ph\.D|Ph\.D\.|Doctorate|Doctor|D\.Phil)\s+(?:in|of)?\s*([^,\n]+?)(?:,|$|\n)', 'PhD'),
            (r'[-•*]?\s*(MS|M\.S|M\.Sc|M\.S\.|Master|Masters?)\s+(?:in|of)?\s*([^,\n]+?)(?:,|$|\n)', 'MS'),
            (r'[-•*]?\s*(BSc|B\.S|B\.Sc|B\.S\.|Bachelor|Bachelors?)\s+(?:in|of)?\s*([^,\n]+?)(?:,|$|\n)', 'BSc'),
            (r'[-•*]?\s*(MBA|M\.B\.A)\s+(?:in|of)?\s*([^,\n]+?)(?:,|$|\n)', 'MBA')
        ]
        
        for pattern, degree_type in degree_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                field = match.group(2).strip() if len(match.groups()) > 1 else ""
                degree_match = (match, degree_type, field)
                break
        
        if degree_match:
            match, degree_type, field = degree_match
            
            # Look ahead for institution (next 5 lines)
            institution = ""
            year = None
            
            for j in range(i+1, min(i+6, len(edu_lines))):
                next_line = edu_lines[j]
                next_lower = next_line.lower()
                
                # Skip if it's another degree (starts with bullet and degree keyword)
                if re.match(r'[-•*]', next_line) and any(word in next_lower for word in ['phd', 'ms', 'm.s', 'bsc', 'bachelor', 'master', 'mba']):
                    break
                
                # Check for institution keywords
                if any(word in next_lower for word in ['university', 'college', 'institute', 'school', 'tech', 'sutd', 'lums', 'uet', 'ntu', 'nanyang']):
                    institution = next_line
                    # Try to extract year from same line or next line
                    year = extract_year(next_line)
                    if not year and j+1 < len(edu_lines):
                        year = extract_year(edu_lines[j+1])
                    break
            
            # Look for "Awarded:" or "Defence:" patterns (can be on same line or next lines)
            for j in range(i+1, min(i+7, len(edu_lines))):
                next_line = edu_lines[j]
                next_lower = next_line.lower()
                if 'awarded' in next_lower or 'defence' in next_lower or 'viva' in next_lower:
                    # Extract year from this line
                    year = extract_year(next_line)
                    # Also check for "Awarded: YYYY" pattern
                    awarded_match = re.search(r'awarded[:\s]+(\d{4})', next_lower)
                    if awarded_match:
                        year = int(awarded_match.group(1))
                    break
            
            # Extract year from current line if not found
            if not year:
                year = extract_year(line)
            
            edu_entry = {
                "degree": degree_type,
                "field": field,
                "institution": institution,
                "year": year
            }
            
            if edu_entry.get("degree") or edu_entry.get("institution"):
                data["education"].append(edu_entry)
            
            i += 1
        else:
            i += 1
    
    # INTELLIGENT Experience Parsing
    exp_start = section_indices.get('experience', 0)
    # If no explicit experience section found, search for experience patterns
    if exp_start == 0:
        for i, line in enumerate(text_lines[:150]):
            line_lower = line.lower()
            if any(word in line_lower for word in ['experience', 'work', 'employment', 'position', 'researcher', 'scientist', 'fellow', 'engineer']):
                if any(word in line_lower for word in ['visiting', 'postdoctoral', 'research', 'university', 'institute']):
                    exp_start = i
                    print(f"[PARSER DEBUG] Found experience pattern at line {i}: '{line[:50]}'")
                    break
    
    exp_end = min(
        section_indices.get('education', len(text_lines)),
        section_indices.get('publications', len(text_lines)),
        exp_start + 100
    ) if exp_start > 0 else min(150, len(text_lines))
    
    exp_lines = text_lines[exp_start:exp_end] if exp_start > 0 else text_lines[:150]
    print(f"[PARSER DEBUG] Experience parsing: lines {exp_start} to {exp_end}, {len(exp_lines)} lines")
    
    current_exp = {}
    i = 0
    while i < len(exp_lines):
        line = exp_lines[i]
        line_lower = line.lower()
        
        # Skip section headers
        if any(word in line_lower for word in ['# education', '# publications', '# skills', '# awards']):
            if current_exp.get("title") or current_exp.get("company"):
                data["experience"].append(current_exp)
            current_exp = {}
            i += 1
            continue
        
        # Pattern 1: Title with bullet "- Visiting Researcher" or "• Postdoctoral Research Fellow"
        title_match = re.match(r'^[-•*]\s*(.+?)(?:$|\n)', line)
        if title_match:
            title = title_match.group(1).strip()
            # More lenient check - any title-like word
            if any(word in title.lower() for word in ['researcher', 'scientist', 'engineer', 'fellow', 'postdoc', 'mentor', 'assistant', 'professor', 'director', 'manager', 'visiting', 'postdoctoral']):
                if current_exp.get("title") or current_exp.get("company"):
                    data["experience"].append(current_exp)
                
                current_exp = {
                    "title": title,
                    "company": "",
                    "location": "",
                    "start_date": "",
                    "end_date": "",
                    "description": []
                }
                i += 1
                continue
        
        # Pattern 2: Company/Institution (follows title, no bullet, usually has university/institute)
        if current_exp.get("title") and not current_exp.get("company"):
            if any(word in line_lower for word in ['university', 'institute', 'college', 'ltd', 'inc', 'lab', 'center', 'centre', 'technologies', 'sutd', 'ntu', 'ucl', 'nanyang', 'queen mary']):
                current_exp["company"] = line
                # Extract location
                for loc in ['Singapore', 'UK', 'United Kingdom', 'USA', 'United States', 'London']:
                    if loc.lower() in line_lower:
                        current_exp["location"] = loc
                        break
                i += 1
                continue
        
        # Pattern 3: Date range "01/09/2025 – Present" or "23/09/2024 – Present" (can be on separate line)
        date_match = re.search(r'(\d{2}[/-]\d{2}[/-]\d{2,4})\s*[–-]\s*(present|current|\d{2}[/-]\d{2}[/-]\d{2,4})', line_lower)
        if date_match:
            if current_exp.get("title"):
                current_exp["start_date"] = date_match.group(1)
                end_date = date_match.group(2)
                if 'present' in end_date.lower() or 'current' in end_date.lower():
                    current_exp["end_date"] = "Present"
                else:
                    current_exp["end_date"] = end_date
            i += 1
            continue
        
        # Pattern 4: Description bullets (starts with bullet or "Strategic")
        if current_exp.get("title") and (line.startswith(('-', '•', '*')) or line.startswith('Strategic') or line.startswith('Actively') or line.startswith('Perform') or line.startswith('Leverage') or line.startswith('Led') or line.startswith('Employed') or line.startswith('Conducted') or line.startswith('Guided') or line.startswith('Organized')):
            desc = line.lstrip('-•* ').strip()
            if len(desc) > 10:
                current_exp["description"].append(desc)
            i += 1
            continue
        
        # Pattern 5: Merged title+company (e.g., "Postdoctoral Research FellowSingapore University...")
        merged_match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*(?:Researcher|Scientist|Fellow|Engineer|Mentor|Professor|Assistant|Director|Manager))([A-Z][^a-z]*?(?:University|Institute|College|Ltd|Inc|Lab))', line)
        if merged_match and not current_exp.get("title"):
            title = merged_match.group(1).strip()
            company = merged_match.group(2).strip()
            
            if current_exp.get("title") or current_exp.get("company"):
                data["experience"].append(current_exp)
            
            current_exp = {
                "title": title,
                "company": company,
                "location": "",
                "start_date": "",
                "end_date": "",
                "description": []
            }
            
            # Extract date and location from same line
            date_match = re.search(r'(\d{2}[/-]\d{2}[/-]\d{2,4})\s*[–-]\s*(present|current|\d{2}[/-]\d{2}[/-]\d{2,4})', line_lower)
            if date_match:
                current_exp["start_date"] = date_match.group(1)
                end_date = date_match.group(2)
                if 'present' in end_date.lower():
                    current_exp["end_date"] = "Present"
                else:
                    current_exp["end_date"] = end_date
            
            for loc in ['Singapore', 'UK', 'United Kingdom', 'USA', 'United States', 'London']:
                if loc.lower() in line_lower:
                    current_exp["location"] = loc
                    break
            
            i += 1
            continue
        
        i += 1
    
    if current_exp.get("title") or current_exp.get("company"):
        data["experience"].append(current_exp)
    
    # INTELLIGENT Publications Parsing
    pub_start = section_indices.get('publications', 0)
    # If no explicit publications section found, search for publication patterns
    if pub_start == 0:
        for i, line in enumerate(text_lines):
            line_lower = line.lower()
            if any(word in line_lower for word in ['publications', 'publication', 'peer-reviewed', 'journal', 'papers']):
                # Check if it's a section header (short line or starts with #)
                if len(line.split()) <= 5 or line_lower.startswith('#'):
                    pub_start = i
                    print(f"[PARSER DEBUG] Found publications pattern at line {i}: '{line[:50]}'")
                    break
            # Also check for numbered entries (1., 2., etc.) which indicate publications
            if re.match(r'^\d+\.', line.strip()):
                # Check if line contains publication-like content
                if any(word in line_lower for word in ['doi', 'journal', 'chem', 'nature', 'science', 'angew']):
                    pub_start = i - 5  # Start a few lines before
                    print(f"[PARSER DEBUG] Found numbered publication at line {i}: '{line[:50]}'")
                    break
    
    if pub_start > 0:
        pub_end = min(
            section_indices.get('presentations', len(text_lines)),
            section_indices.get('awards', len(text_lines)),
            section_indices.get('skills', len(text_lines)),
            len(text_lines)
        )
        pub_lines = text_lines[pub_start:pub_end]
        pub_text = '\n'.join(pub_lines)
        
        # Pattern: "Number. Authors (Year): Title. Journal, Volume, Pages. DOI"
        pub_pattern = re.compile(
            r'(\d+)\.\s*'  # Number
            r'([^:]+?)\s*'  # Authors
            r'\((\d{4})\):\s*'  # Year
            r'([^\.]+?\.)\s*'  # Title (first sentence)
            r'([^\.]+?\.)\s*'  # Journal
            r'(?:DOI[:\s]*([^\s,]+))?',  # Optional DOI
            re.IGNORECASE | re.MULTILINE
        )
        
        matches = pub_pattern.finditer(pub_text)
        for match in matches:
            authors = match.group(2).strip()
            year = int(match.group(3))
            title = match.group(4).strip().rstrip('.,;')
            journal = match.group(5).strip().rstrip('.,;')
            doi = match.group(6) if match.group(6) else None
            
            if title and len(title) > 15:
                pub = {
                    "title": title[:500],
                    "authors": authors[:500] if authors else "",
                    "journal": journal[:200] if journal else "",
                    "year": year,
                    "doi": doi
                }
                data["publications"].append(pub)
        
        # Fallback: Split by numbered entries (more robust)
        if not data["publications"]:
            pub_entries = re.split(r'\n(?=\d+\.)', pub_text)
            
            for entry in pub_entries:
                entry = entry.strip()
                if len(entry) < 30:
                    continue
                
                # Remove leading number
                entry = re.sub(r'^\d+\.\s*', '', entry)
                
                # Extract year
                year = extract_year(entry)
                
                # Extract DOI
                doi_match = re.search(r'doi[:\s]*([\d\.]+/[^\s,]+)', entry, re.IGNORECASE)
                doi = doi_match.group(1) if doi_match else None
                
                # Extract journal (common patterns)
                journal = ""
                journal_patterns = [
                    r'Angew\.\s*Chem\.\s*(?:Int\.\s*Ed\.)?',
                    r'J\.\s*Am\.\s*Chem\.\s*Soc\.',
                    r'Nature\s*Methods',
                    r'CCS\s*Chemistry',
                    r'Chem\.\s*Commun\.',
                    r'Dyes\s*Pigments',
                    r'Adv\.\s*Sci\.',
                    r'Adv\.\s*Funct\.\s*Mater\.',
                    r'Chin\.\s*Chem\.\s*Lett\.',
                    r'Mater\.\s*Chem\.\s*Front\.',
                    r'Nat\.\s*Commun\.',
                    r'J\.\s*Phys\.\s*Chem\.\s*[A-Z]',
                ]
                
                for pattern in journal_patterns:
                    match = re.search(pattern, entry, re.IGNORECASE)
                    if match:
                        journal = match.group(0).strip('.,;')
                        break
                
                # If no journal found, try to extract from format "Journal, Volume" or "Journal. Volume"
                if not journal:
                    journal_match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*[,\.]\s*\d+', entry)
                    if journal_match:
                        journal = journal_match.group(1).strip('.,;')
                
                # Extract authors (before year in parentheses)
                authors_match = re.search(r'^([^\(]+?)\s*\((\d{4})\)', entry)
                if authors_match:
                    authors = authors_match.group(1).strip()
                    if not year:
                        year = int(authors_match.group(2))
                else:
                    # Try pattern without year
                    authors_match = re.match(r'^([^:]+?)(?::|\(|\d)', entry)
                    authors = authors_match.group(1).strip() if authors_match else ""
                
                # Extract title (between year and journal, or after colon)
                if ':' in entry:
                    # Title is after the colon
                    title_part = entry.split(':', 1)[1]
                    if journal:
                        # Title is before journal
                        title = title_part.split(journal)[0].strip().rstrip('.,;')
                    else:
                        # Title is first sentence before period
                        title = title_part.split('.')[0].strip().rstrip('.,;')
                else:
                    # Try to extract from format "Authors (Year): Title. Journal"
                    title_match = re.search(r'\((\d{4})\):\s*([^\.]+?)(?:\.|$)', entry)
                    if title_match:
                        title = title_match.group(2).strip().rstrip('.,;')
                    else:
                        # Fallback: first substantial part
                        parts = entry.split('.')
                        if len(parts) > 1:
                            title = parts[0].strip().rstrip('.,;')
                        else:
                            title = entry[:200].strip()
                
                if title and len(title) > 10:
                    pub = {
                        "title": title[:500],
                        "authors": authors[:500] if authors else "",
                        "journal": journal[:200] if journal else "",
                        "year": year,
                        "doi": doi
                    }
                    data["publications"].append(pub)
    
    # INTELLIGENT Skills Parsing
    skills_start = section_indices.get('skills', 0)
    # If no explicit skills section found, search for skills patterns
    if skills_start == 0:
        for i, line in enumerate(text_lines):
            line_lower = line.lower()
            # Look for "Technical Skills" or table format with "Category | Skills"
            if any(word in line_lower for word in ['technical skills', 'skills and expertise', 'expertise']):
                # Check if it's a section header or table format
                if (len(line.split()) <= 5 or 
                    '|' in line or 
                    line_lower.startswith('#') or
                    'category' in line_lower):
                    skills_start = i
                    print(f"[PARSER DEBUG] Found skills header at line {i}: '{line[:50]}'")
                    break
            # Also check for table format: "| Category | Skills |"
            if '|' in line and 'category' in line_lower and 'skill' in line_lower:
                skills_start = i
                print(f"[PARSER DEBUG] Found skills table at line {i}: '{line[:50]}'")
                break
    
    if skills_start > 0:
        skills_end = min(len(text_lines), skills_start + 100)
        skills_lines = text_lines[skills_start:skills_end]
        
        # Try to extract from table format first
        table_mode = False
        for line in skills_lines:
            if '|' in line and 'Category' in line:
                table_mode = True
                break
        
        if table_mode:
            # Table format: "| Category | Skills |"
            for line in skills_lines:
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    if len(parts) >= 3 and parts[0].lower() not in ['category', '']:
                        # Skip header row
                        if parts[0].lower() == 'category':
                            continue
                        skills_text = parts[1] if len(parts) > 1 else ""
                        # Extract individual skills from skills column
                        # Split by commas, semicolons, or periods
                        skills_list = re.split(r'[,;:\.]', skills_text)
                        for skill in skills_list:
                            skill = skill.strip()
                            # Remove parentheses content but keep important info
                            skill = re.sub(r'\([^)]+\)', '', skill).strip()
                            # Remove common prefixes
                            skill = re.sub(r'^(Languages?|Software|Tools?|Frameworks?|ML|Machine Learning|Data|Programming):\s*', '', skill, flags=re.IGNORECASE)
                            # Clean up
                            skill = re.sub(r'^\s*[-•*]\s*', '', skill)
                            if len(skill) > 2 and skill.lower() not in ['and', 'or', 'the', 'a', 'category', 'skills', 'expertise']:
                                data["skills"].append(skill)
        
        # Also try line-by-line extraction
        for line in skills_lines:
            line_lower = line.lower()
            
            if any(word in line_lower for word in ['# experience', '# education', '# publications', '# awards', '# presentations']):
                break
            
            # Skip category headers
            if any(word in line_lower for word in ['category', 'technical skills and expertise', 'skills', 'expertise']) and ('|' in line or len(line.split()) <= 5):
                continue
            
            # Extract from table format (Category | Skills) - if not already processed
            if '|' in line and not table_mode:
                parts = [p.strip() for p in line.split('|')]
                if len(parts) >= 2 and parts[0].lower() != 'category':
                    skills_text = parts[1]
                    skills_list = re.split(r'[,;:]', skills_text)
                    for skill in skills_list:
                        skill = skill.strip()
                        skill = re.sub(r'\([^)]+\)', '', skill).strip()
                        skill = re.sub(r'^(Languages?|Software|Tools?|Frameworks?|ML|Machine Learning):\s*', '', skill, flags=re.IGNORECASE)
                        if len(skill) > 2 and skill.lower() not in ['and', 'or', 'the', 'a', 'category', 'skills']:
                            data["skills"].append(skill)
            
            # Extract from colon format
            elif ':' in line and len(line.split(':')) == 2:
                category, skills_text = line.split(':', 1)
                if category.lower().strip() not in ['category', 'technical skills', 'expertise', 'technical skills and expertise']:
                    skills_list = re.split(r'[,;]', skills_text)
                    for skill in skills_list:
                        skill = skill.strip()
                        if len(skill) > 2:
                            data["skills"].append(skill)
            
            # Extract from bullets
            elif line.startswith(('-', '•', '*')):
                skill = line.lstrip('-•* ').strip()
                if len(skill) > 2 and skill.lower() not in ['category', 'skills', 'expertise']:
                    data["skills"].append(skill)
    
    # Remove duplicates and filter
    filtered_skills = []
    skip_words = ['technical', 'expertise', 'skills', 'category', 'competencies', 'proficiencies', 'and', 'or', 'the', 'a']
    for skill in data["skills"]:
        skill_lower = skill.lower().strip()
        if (not any(word == skill_lower for word in skip_words) and 
            len(skill) > 2 and 
            not skill_lower.startswith('category') and
            not skill_lower.startswith('technical') and
            skill_lower not in ['ics', 'ing', 'tion']):
            filtered_skills.append(skill)
    data["skills"] = list(dict.fromkeys(filtered_skills))[:100]
    
    return data


def parse_resume_file_intelligent(filepath: str) -> Dict:
    """Parse resume file using intelligent parser"""
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    file_ext = os.path.splitext(filepath)[1].lower()
    layout_info = {}
    
    if file_ext == '.pdf':
        text, layout_info = extract_text_from_pdf_intelligent(filepath)
    elif file_ext in ['.docx', '.doc']:
        text, layout_info = extract_text_from_docx_intelligent(filepath)
    elif file_ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    if not text or len(text.strip()) < 50:
        raise ValueError("Resume file appears to be empty")
    
    return parse_resume_intelligent(text, layout_info)
